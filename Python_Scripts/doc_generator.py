import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import (
    WD_PARAGRAPH_ALIGNMENT
)


def add_bold_text(paragraph, text):

    """
    Convert markdown bold and italics
    into Word formatting.
    """

    parts = re.split(
        r"(\*\*.*?\*\*)",
        text
    )

    for part in parts:

        if (
            part.startswith("**")
            and part.endswith("**")
        ):

            run = paragraph.add_run(
                part[2:-2]
            )

            run.bold = True

        else:

            # Remove markdown italics
            cleaned = re.sub(
                r"\*(.*?)\*",
                r"\1",
                part
            )

            paragraph.add_run(
                cleaned
            )


def markdown_to_docx(
    content,
    output_path
):

    doc = Document()

    # ==================================
    # NORMAL FONT
    # ==================================

    style = doc.styles["Normal"]

    style.font.name = "Calibri"

    style.font.size = Pt(11)

    # ==================================
    # PROCESS MARKDOWN
    # ==================================

    for line in content.split("\n"):

        line = line.strip()

        if not line:
            continue

        # Ignore separators
        if line == "---":
            continue

        # ==================================
        # H1
        # ==================================

        if line.startswith("# "):

            heading = doc.add_heading(
                level=1
            )

            heading.alignment = (
                WD_PARAGRAPH_ALIGNMENT.LEFT
            )

            run = heading.add_run(
                line.replace("# ", "")
            )

            run.bold = True
            run.font.size = Pt(20)

        # ==================================
        # H2
        # ==================================

        elif line.startswith("## "):

            heading = doc.add_heading(
                level=2
            )

            run = heading.add_run(
                line.replace("## ", "")
            )

            run.bold = True
            run.font.size = Pt(16)

        # ==================================
        # H3
        # ==================================

        elif line.startswith("### "):

            heading = doc.add_heading(
                level=3
            )

            run = heading.add_run(
                line.replace("### ", "")
            )

            run.bold = True
            run.font.size = Pt(14)

        # ==================================
        # H4
        # ==================================

        elif line.startswith("#### "):

            heading = doc.add_heading(
                level=4
            )

            run = heading.add_run(
                line.replace("#### ", "")
            )

            run.bold = True
            run.font.size = Pt(12)

        # ==================================
        # BULLETS (-)
        # ==================================

        elif line.startswith("- "):

            paragraph = doc.add_paragraph(
                style="List Bullet"
            )

            add_bold_text(
                paragraph,
                line[2:]
            )

        # ==================================
        # BULLETS (*)
        # ==================================

        elif line.startswith("* "):

            paragraph = doc.add_paragraph(
                style="List Bullet"
            )

            add_bold_text(
                paragraph,
                line[2:]
            )

        # ==================================
        # NORMAL TEXT
        # ==================================

        else:

            # Remove entire line italics
            if (
                line.startswith("*")
                and line.endswith("*")
                and len(line) > 2
            ):
                line = line[1:-1]

            paragraph = doc.add_paragraph()

            add_bold_text(
                paragraph,
                line
            )

    # ==================================
    # SAVE
    # ==================================

    doc.save(
        output_path
    )